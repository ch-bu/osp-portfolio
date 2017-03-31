#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
This files converts multiple word documents for the osp-internship
into a spreadsheet.

Compressed folders are uncompressed and every word-file is
read for its contents. The content is then converted to
a csv file which is then converted to a excel spreadsheet.
"""

# http://stackoverflow.com/questions/41550620/python-docx-get-info-from-dropdownlist-in-table

import sys
import zipfile
import tarfile
import re
import fnmatch
# import glob
import glob2
import os
import shutil
from docx import Document
from wordhelper import get_data
from xlsxwriter.workbook import Workbook
import json
import csv

# Magic happens
# note the unicode errors are gone
reload(sys)
sys.setdefaultencoding('utf8')


def unzip_files(compressed_file):

    print('Processing %s' % compressed_file)

    # Decompress zip file
    if compressed_file.endswith('.zip'):
        # Unzip file
        zip_ref = zipfile.ZipFile(compressed_file, 'r')
        zip_ref.extractall(temp_path)

        # Analyze data
        subject_results = analyze_files(temp_path)

        # Close zip file
        zip_ref.close()
    # Decompress tar file
    elif compressed_file.endswith('.tar'):
        tar = tarfile.open(compressed_file, "r:")
        tar.extractall(temp_path)

        # Analyze data
        subject_results = analyze_files(temp_path)

        # Close tar file
        tar.close()

    # # Delete directory for new subject
    shutil.rmtree(temp_path)

    # # Return results
    return subject_results


def return_file_content(document, word_file):

    # Store every result in dictionary that
    # can be processed later
    results = {}

    # Add title of document to results
    title = document.paragraphs[0].text

    # Make sure not to include the first paragraph
    # that includes the title
    iter_paragraphs = iter(document.paragraphs)
    next(iter_paragraphs)

    # Get data from current word file
    results[title] = get_data(title, iter_paragraphs, document, word_file)

    # Return content of word file in dictionary
    return results


def analyze_files(temp_path):
    # Get all docx files
    word_files = [os.path.join(root, name) for root, dirs, files \
                    in os.walk(temp_path)
                    for name in files if (not re.match(r'.*~.*', name)) and \
                    (not re.match(r'.*MACOSX.*', root)) and \
                    name.endswith('.docx') and (not re.match(r'\.\_.*', name))]

    res = []

    # Loop over every word file
    for word_file in word_files:
        # print(word_file)
        # Get word document
        document = Document(word_file)

        res.append(return_file_content(document, word_file))

    return res


# Run code only if the program is run by itself and not
# when it is imported from another module
if __name__ == '__main__':

    # Get dir of portfolio folder two
    file_path = os.path.dirname(os.path.realpath(__file__))
    portfolio_path = file_path + '/portfolios-two'
    temp_path = file_path + '/temp'

    # Get all compressed files and save path in zip_files
    # file_extensions = ['.zip', '.tar', '.7s']
    file_extensions = ['.zip', '.tar']
    zip_files = [portfolio[0] + '/' + file for portfolio \
                in os.walk(portfolio_path) for file in portfolio[2] \
                if file.endswith(tuple(file_extensions))]

    # Get data from all subjects
    data_subjects = list(map(lambda compressed_file: \
            unzip_files(compressed_file), zip_files))

    # # Write results to disk
    with open('results/portfolio-two.csv', 'w') as f:

        # Create writer for csv file
        writer = csv.writer(f, delimiter=';', quotechar='"', quoting=csv.QUOTE_MINIMAL, lineterminator='\n')

        # Write header
        writer.writerow(['Vorname', 'Nachname', 'Matrikelnummer',
                        'Mail', 'Hauptfach.1', 'Hauptfach.2',
                        'Begleitung Alltag Lehrperson',
                        'beobachten.one.activity', 'beobachten.one.content',
                        'beobachten.two.activity', 'beobachten.two.content',
                        'beobachten.three.activity', 'beobachten.three.content',
                        'Wahlbeobachtung',
                        'Erste Durchführung einer zentralen Tätigkeit',
                        'Zweite Durchführung einer zentralen Tätigkeit',
                        'Schlüsselsituation',
                        'Interview mit einer Lehrkraft',
                        'Wahl-Aufgabe'])

        # Loop over every subject
        for subject in data_subjects:

            my_dict = {}

            # Loop over every file
            for word_file in subject:
                # Get key of file
                key = list(word_file.keys())[0]

                # Add personal information
                if key == 'Zu Ihrer Person':
                    curr_dict = word_file[key]
                    my_dict['Vorname'] = curr_dict['Vorname']
                    my_dict['Nachname'] = curr_dict['Nachname']
                    my_dict['Matrikelnummer'] = curr_dict['Matrikelnummer']
                    my_dict['Mail'] = curr_dict['Mail']
                    my_dict['Hauptfach.1'] = curr_dict['Hauptfach_1']
                    my_dict['Hauptfach.2'] = curr_dict['Hauptfach_2']
                    # my_dict = {**word_file[key]}
                if key == 'Beobachten. Dritte Tätigkeit':
                    curr_dict = word_file[key]
                    my_dict['beobachten.three.activity'] = curr_dict['activity']
                    my_dict['beobachten.three.content'] = curr_dict['content'].replace('\n', ' ').replace('\r', '')
                elif key == 'Beobachten. Zweite Tätigkeit':
                    curr_dict = word_file[key]
                    my_dict['beobachten.two.activity'] = curr_dict['activity']
                    my_dict['beobachten.two.content'] = curr_dict['content'].replace('\n', ' ').replace('\r', '')
                elif key == 'Beobachten. Erste Tätigkeit':
                    curr_dict = word_file[key]
                    my_dict['beobachten.one.activity'] = curr_dict['activity']
                    my_dict['beobachten.one.content'] = curr_dict['content'].replace('\n', ' ').replace('\r', '')
                elif key == 'Wahlbeobachtung' or \
                     key == 'Erste Durchführung einer zentralen Tätigkeit' or \
                     key == 'Zweite Durchführung einer zentralen Tätigkeit' or \
                     key == 'Schlüsselsituation' or \
                     key == 'Begleitung Alltag Lehrperson' or \
                     key == 'Interview mit einer Lehrkraft' or \
                     key == 'Wahl-Aufgabe':
                        # Store content temporarily
                        content = word_file[key].replace('\n', ' ').replace('\r', '')

                        # Escape special characters to avoid stupid bug
                        # Not beautiful, but it works.
                        if key == 'Erste Durchführung einer zentralen Tätigkeit':
                            key = 'Erste Durchfuehrung einer zentralen Taetigkeit'
                        elif key == 'Zweite Durchführung einer zentralen Tätigkeit':
                            key = 'Zweite Durchfuehrung einer zentralen Taetigkeit'
                        elif key == 'Schlüsselsituation':
                            key = 'Schluesselsituation'

                        my_dict[key] = content

            # print(my_dict.get('Erste Durchführung einer zentralen Tätigkeit', ''))

            # Write row for subject
            writer.writerow([my_dict.get('Vorname', ''), my_dict.get('Nachname', ''), my_dict.get('Matrikelnummer', ''),
                            my_dict.get('Mail', ''), my_dict.get('Hauptfach.1', ''), my_dict.get('Hauptfach.2', ''),
                            my_dict.get('Begleitung Alltag Lehrperson', ''),
                            my_dict.get('beobachten.one.activity', ''), my_dict.get('beobachten.one.content', ''),
                            my_dict.get('beobachten.two.activity', ''), my_dict.get('beobachten.two.content', ''),
                            my_dict.get('beobachten.three.activity', ''), my_dict.get('beobachten.three.content', ''),
                            my_dict.get('Wahlbeobachtung', ''),
                            my_dict.get('Erste Durchfuehrung einer zentralen Taetigkeit', ''),
                            my_dict.get('Zweite Durchfuehrung einer zentralen Taetigkeit', ''),
                            my_dict.get('Schluesselsituation', ''),
                            my_dict.get('Interview mit einer Lehrkraft', ''),
                            my_dict.get('Wahl-Aufgabe', '')])

    # Write Excel spreadsheet
    # for csvfile in glob.glob(os.path.join('results', '*.csv')):
    workbook = Workbook('results/portfolio-two.xlsx')
    worksheet = workbook.add_worksheet()

    with open('results/portfolio-two.csv', 'rt') as f:
        reader = csv.reader(f, delimiter=';')
        for r, row in enumerate(reader):
            for c, col in enumerate(row):
                worksheet.write(r, c, col)

    workbook.close()
