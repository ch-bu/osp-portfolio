#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
This files converts multiple word documents for the osp-internship
into a spreadsheet.

Compressed folders are uncompressed and every word-file is
read for its contents. The content is then converted to
a csv file which is then converted to a excel spreadsheet.
"""

# http://stackoverflow.com/questions/41550620/python-docx-get-info-from-dropdownlist-in-table

import sys
import zipfile
import tarfile
import re
import fnmatch
# import glob
import glob2
import os
import shutil
from docx import Document
from wordhelper import get_data
from xlsxwriter.workbook import Workbook
import json
import csv

# Magic happens
# note the unicode errors are gone
reload(sys)
sys.setdefaultencoding('utf8')


def unzip_files(compressed_file):

    # print('Processing %s' % compressed_file)

    # Decompress zip file
    if compressed_file.endswith('.zip'):
        # Unzip file
        zip_ref = zipfile.ZipFile(compressed_file, 'r')
        zip_ref.extractall(temp_path)

        # Analyze data
        subject_results = analyze_files(temp_path)

        try:
            person = subject_results[0]['Zu Ihrer Person']
        except KeyError:
            print('### Key Error')
            print(compressed_file)
        except IndexError:
            print('### Index Error')
            print(compressed_file)

        if subject_results[0]['Zu Ihrer Person'] == '':
            print('### Empty file')
            print(compressed_file)

        # Close zip file
        zip_ref.close()

    # Delete directory for new subject
    shutil.rmtree(temp_path)

    # Return results
    return subject_results
    # return None


def return_file_content(document, word_file):

    # Store every result in dictionary that
    # can be processed later
    results = {}

    # Add title of document to results
    title = document.paragraphs[0].text

    # Make sure not to include the first paragraph
    # that includes the title
    iter_paragraphs = iter(document.paragraphs)
    next(iter_paragraphs)

    # Get data from current word file
    results[title] = get_data(title, iter_paragraphs, document, word_file)

    # Return content of word file in dictionary
    return results


def get_data(title, iter_paragraphs, document, word_file):
    """Extracts the data from a word file from an
    specific file
    """

    # Check which word file is analyzed and behave
    # appropriately
    if title == 'Zu Ihrer Person':
        return get_person(get_simple_data(iter_paragraphs))
    elif title == 'Reflexionsaufgabe Berufswahl' or \
         title == 'Reflektionsaufgabe Kompetenzentwicklung':
        return get_simple_data(iter_paragraphs)
    # else:
    #     return None

def get_simple_data(paragraphs):
    """Get data from paragraphs
    """

    # Init result string
    result = ''

    # Loop over every paragraph
    for paragraph in paragraphs:
        # Append new paragraph to previous paragraph
        result = result + ' ' + paragraph.text

    # Return string of paragraphs
    return result.replace('\n', ' ').replace('\r', '')


def get_person(person_string):
    """Extracts the data about
    a specific person
    """

    # Gets the values of the specific person
    first_name = re.search('Vorname:(.+?)Nachname', person_string).group(1).strip()
    last_name = re.search('Nachname:(.+?)Matrikelnummer', person_string).group(1).strip()
    number = re.search('Matrikelnummer:(.+?)E-Mail', person_string).group(1).strip()
    mail = re.search('E-Mail:(.+?)Hauptfach', person_string).group(1).strip()
    subject_one = re.search('Hauptfach 1:(.+?)Hauptfach', person_string).group(1).strip()
    subject_two = re.search('Hauptfach 2:(.+?)$', person_string).group(1).strip()

    result = {'Vorname': first_name, 'Nachname': last_name, 'Matrikelnummer': number, \
        'Mail': mail, 'Hauptfach_1': subject_one, 'Hauptfach_2': subject_two}

    return result

def analyze_files(temp_path):
    # Get all docx files
    word_files = [os.path.join(root, name) for root, dirs, files \
                    in os.walk(temp_path)
                    for name in files if (not re.match(r'.*~.*', name)) and \
                    (not re.match(r'.*MACOSX.*', root)) and \
                    name.endswith('.docx') and (not re.match(r'\.\_.*', name))]

    res = []

    # Loop over every word file
    for word_file in word_files:
        # Get word document
        document = Document(word_file)

        res.append(return_file_content(document, word_file))


    # print(res[0]['Zu Ihrer Person'])

    return res


# Run code only if the program is run by itself and not
# when it is imported from another module
if __name__ == '__main__':

    # Get dir of portfolio folder two
    file_path = os.path.dirname(os.path.realpath(__file__))
    portfolio_path = file_path + '/portfolios-three'
    temp_path = file_path + '/temp'

    # Get all compressed files and save path in zip_files
    # file_extensions = ['.zip', '.tar', '.7s']
    file_extensions = ['.zip']
    zip_files = [portfolio[0] + '/' + file for portfolio \
                in os.walk(portfolio_path) for file in portfolio[2] \
                if file.endswith(tuple(file_extensions))]

    # Get data from all subjects
    data_subjects = list(map(lambda compressed_file: \
            unzip_files(compressed_file), zip_files))

    # print(data_subjects)

    # Write results to disk
    with open('results/portfolio-three.csv', 'w') as f:

        # Create writer for csv file
        writer = csv.writer(f, delimiter=';', quotechar='"', quoting=csv.QUOTE_MINIMAL, lineterminator='\n')

        # Write header
        writer.writerow(['Vorname', 'Nachname', 'Matrikelnummer',
                        'Mail', 'Hauptfach.1', 'Hauptfach.2',
                        'Reflexionsaufgabe Berufswahl',
                        'Reflektionsaufgabe Kompetenzentwicklung'])

        # Loop over every subject
        for subject in data_subjects:

            my_dict = {}

            # Loop over every file
            for word_file in subject:
                # Get key of file
                key = list(word_file.keys())[0]

                # Add personal information
                if key == 'Zu Ihrer Person':
                    curr_dict = word_file[key]
                    my_dict['Vorname'] = curr_dict['Vorname']
                    my_dict['Nachname'] = curr_dict['Nachname']
                    my_dict['Matrikelnummer'] = curr_dict['Matrikelnummer']
                    my_dict['Mail'] = curr_dict['Mail']
                    my_dict['Hauptfach.1'] = curr_dict['Hauptfach_1']
                    my_dict['Hauptfach.2'] = curr_dict['Hauptfach_2']
                elif key == 'Reflexionsaufgabe Berufswahl' or \
                     key == 'Reflektionsaufgabe Kompetenzentwicklung':
                        # Store content temporarily
                        content = word_file[key].replace('\n', ' ').replace('\r', '')

                        my_dict[key] = content

            # print(my_dict.get('Erste Durchführung einer zentralen Tätigkeit', ''))

            # Write row for subject
            writer.writerow([my_dict.get('Vorname', ''), my_dict.get('Nachname', ''), my_dict.get('Matrikelnummer', ''),
                            my_dict.get('Mail', ''), my_dict.get('Hauptfach.1', ''), my_dict.get('Hauptfach.2', ''),
                            my_dict.get('Reflexionsaufgabe Berufswahl', ''),
                            my_dict.get('Reflektionsaufgabe Kompetenzentwicklung', '')])

    # Write Excel spreadsheet
    # for csvfile in glob.glob(os.path.join('results', '*.csv')):
    workbook = Workbook('results/portfolio-three.xlsx')
    worksheet = workbook.add_worksheet()

    with open('results/portfolio-three.csv', 'rt') as f:
        reader = csv.reader(f, delimiter=';')
        for r, row in enumerate(reader):
            for c, col in enumerate(row):
                worksheet.write(r, c, col)

    workbook.close()
